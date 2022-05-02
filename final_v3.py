import openpyxl
from docx import Document
import re

wb = openpyxl.load_workbook(filename = 
'C:\\Users\\shave\\Documents\\Python soubory\\Programy\\Petra excel to word\\NEW2\\BOM.xlsx')
b = wb.worksheets[0]

BOM = list()
for row in range (221,516):
    #print(b[f"A{row}"].value)
    a = dict()
    a["typ"]=b[f"A{row}"].value
    a["oil"]=b[f"B{row}"].value
    a["procent"]=b[f"C{row}"].value
    a["terpen"]=b[f"D{row}"].value
    a["typ_o"]=float(b[f"E{row}"].value)
    a["typ_w"]= format(a["typ_o"]*0.05, "0.3f")
    a["oil_o"]=float(b[f"G{row}"].value)
    a["oil_w"]=format(a["oil_o"]*0.05, "0.3f")
    if b[f"I{row}"].value == "N/A":
        a["terpen_o"]=None
        a["terpen_w"]=None
    else:
        a["terpen_o"]=float(b[f"I{row}"].value)
        a["terpen_w"]=format(a["terpen_o"]*0.05, "0.3f")
    if b[f"K{row}"].value == "N/A" or b[f"K{row}"].value == None:
        a["cbd_crude_o"]=None
        a["cbd_crude_w"]=None
    else:
        a["cbd_crude_o"]=float(b[f"L{row}"].value)
        a["cbd_crude_w"]=format(a["cbd_crude_o"]*0.05, "0.3f")
    if b[f"N{row}"].value == None:
        a["cbg_o"]=None
        a["cbg_w"]=None
    else:
        a["cbg_o"]=float(b[f"O{row}"].value)
        a["cbg_w"]=format(a["cbg_o"]*0.05, "0.3f")
    if b[f"Q{row}"].value == None:
        a["mct_o"]=None
        a["mct_w"]=None
    else:
        a["mct_o"]=float(b[f"R{row}"].value)
        a["mct_w"]=format(a["mct_o"]*0.05, "0.3f")
    BOM.append(a)
#print(BOM)
wb.close

def ztabulky(row):
    wb = openpyxl.load_workbook(filename = 
    'C:\\Users\\shave\\Documents\\Python soubory\\Programy\\Petra excel to word\\NEW2\\DOC2.xlsx')
    c = wb.worksheets[0]
    doc = Document('C:\\Users\\shave\\Documents\\Python soubory\\Programy\\Petra excel to word\\NEW2\\docx2.docx')
    global nazev_soubor
    # Název z DOC2 (B) do doc T0,0,1
    a = doc.tables[0].cell(0,1)
    a.text = c[f"B{row}"].value
    nazev_soubor = c[f"B{row}"].value
    # Product code z DOC2 (A) do doc T0,1,1
    a = doc.tables[0].cell(1,1)
    a.text = c[f"A{row}"].value
    # Starting materials:
    x = c[f"B{row}"].value
    x = x.split()
    print(x)

    #k1 - první klíč
    if "BS" in x and not "(Clear)" in x:
        k1 = "Broad Spectrum"
    if "BS" in x and "(Clear)" in x:
        k1 = "Broad Spectrum Clear"

    #k2 - druhý klíč - oil
    if "Hempseed" in x:
        k2 = "Hempseed Oil"
    elif "Organic" in x and "MCT" in x:
        k2 = "Organic MCT"
    elif not "Organic" in x and "MCT" in x:
        k2 = "MCT Oil"
    elif "Avocado" in x:
        k2 = "Avocado Oil"
    elif "Sunflower" in x:
        k2 ="Sunflower Oil"
    elif "Argan" in x:
        k2 = "Argan Oil"
    elif "Olive" in x:
        k2 = "Olive Oil"
    elif "Grapeseed" in x:
        k2 = "Grapeseed Oil"
    elif "Cod" in x:
        k2 = "Cod Liver Oil"
    elif "Salmon" in x:
        k2 = "Salmon Oil"
    elif "MPG" in x:
        k2 = "MPG"

    #k3 - třetí klíč - procenta
    if "1.67%" in x:
        k3 = 0.0167
    elif "3.33%" in x:
        k3 = 0.0333
    elif "5%" in x:
        k3 = 0.05
    elif "6.67%" in x:
        k3 = 0.0667
    elif "10%" in x:
        k3 = 0.1
    elif "12.5%" in x:
        k3 = 0.125
    elif "15%" in x:
        k3 = 0.15
    elif "20%" in x:
        k3 = 0.2
    elif "25%" in x:
        k3 = 0.25
    elif "30%" in x:
        k3 = 0.3

    #k4 - čtvrtý klíč - terpen/aroma
    if not "with" in x:
        k4 = "No"
    elif "Beef" in x or "Chicken" in x:
        k4 = "Flavour"
    else:
        k4 = "Terpene"

    # nalezení záznamu v dict podle klíčů
    aa = next(i for i in BOM if i["typ"] == k1 and i["oil"]== k2 and i["procent"]== k3 and i["terpen"]== k4)
    print(aa)
    
    # 1 - CBD Isolate 98%
    a = doc.tables[1].cell(1,0)
    a.text = "CBD Isolate 98%"

    a = doc.tables[1].cell(1,1)
    a.text = str(aa["typ_o"])+"g"

    a = doc.tables[1].cell(1,2)
    a.text = str(aa["typ_w"])+"g"
    
    # 2 - vybrat typ oilu
    a = doc.tables[1].cell(2,0)
    y = aa["oil"]
    yy = y.split()
    if not "Oil" in yy:
        y = aa["oil"]+" Oil"
    a.text = y
    
    # 3 - obsah oilu
    a = doc.tables[1].cell(2,1)
    a.text = str(aa["oil_o"])+"g"

    a = doc.tables[1].cell(2,2)
    a.text = str(aa["oil_w"])+"g"
 
    # 4 - Crude oil
    if aa["typ"]=='Broad Spectrum':
        doc.tables[1].add_row()
        a = doc.tables[1].cell(-1,0)
        a.text = "Crude Oil"
        a = doc.tables[1].cell(-1,1)
        a.text = str(aa["cbd_crude_o"])+"g"
        a = doc.tables[1].cell(-1,2)
        a.text = str(aa["cbd_crude_w"])+"g"
       
        # 5 - CBG Isolate
        doc.tables[1].add_row()
        a = doc.tables[1].cell(-1,0)
        a.text = "CBG Isolate 98%"
        a =  doc.tables[1].cell(-1,1)
        a.text = str(aa["cbg_o"])+"g"
        a = doc.tables[1].cell(-1,2)
        a.text = str(aa["cbg_w"])+"g"
    
    # 6 - MCT Oil u BS Salmon a Cod
    if not aa["mct_o"]==None:
        doc.tables[1].add_row()
        a = doc.tables[1].cell(-1,0)
        a.text = "MCT Oil"
        a =  doc.tables[1].cell(-1,1)
        a.text = str(aa["mct_o"])+"g"
        a = doc.tables[1].cell(-1,2)
        a.text = str(aa["mct_w"])+"g"

    # 7 - terpeny a aroma
    if aa["terpen"]=="Terpene":
        doc.tables[1].add_row()
        a = doc.tables[1].cell(-1,0)
        #změnit row 
        x = c[f"B{row}"].value
        xx = re.findall("with (.*)", x)
        a.text = xx[0]+" terpene"
        a =  doc.tables[1].cell(-1,1)
        a.text = str(aa['terpen_o'])+"g"
        a = doc.tables[1].cell(-1,2)
        a.text = str(aa['terpen_w'])+"g"

    if aa["terpen"]=="Flavour":
        doc.tables[1].add_row()
        a = doc.tables[1].cell(-1,0)
        #změnit row 
        x = c[f"B{row}"].value
        xx = re.findall("with (.*)", x)
        a.text = xx[0]+" flavour"
        a =  doc.tables[1].cell(-1,1)
        a.text = str(aa['terpen_o'])+"g"
        a = doc.tables[1].cell(-1,2)
        a.text = str(aa['terpen_w'])+"g"

    # Tabulka 3 - změna CBD%
    a = doc.tables[2].cell(3,1)
    #změnit row
    x = c[f"B{row}"].value
    d = re.findall("^[a-zA-Z\(\)\+ ]+([\d\.]+)%", x)
    a.text = f"{d[0]} %"

    doc.save(f'C:\\Users\\shave\\Documents\\Python soubory\\Programy\\Petra excel to word\\NEW2\\test\\{nazev_soubor}.docx')

def opakovani(x,y):
    for i in range(x,y+1):#x=první řádek v tabulce z které čerpám text, y=poslední řádek
        ztabulky(i)
    wb.close()


opakovani(1228, 1228)